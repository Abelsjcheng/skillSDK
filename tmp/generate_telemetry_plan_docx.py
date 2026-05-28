from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"F:\AIProject\skillSDK")
SOURCE = ROOT / "ai-chat-viewer" / "docs" / "plans" / "2026-05-20-ai-chat-viewer-telemetry-plan.md"
TARGET = ROOT / "ai-chat-viewer" / "docs" / "plans" / "2026-05-20-ai-chat-viewer-telemetry-plan.docx"


@dataclass
class JsonLikeField:
    key: str
    value: str | None = None
    children: list["JsonLikeField"] | None = None
    comment: str | None = None


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def add_heading(document: Document, level: int, text: str) -> None:
    paragraph = document.add_paragraph(style=f"Heading {min(level, 3)}")
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_list_item(document: Document, text: str, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if index < len(lines) - 1:
            run.add_break()
    document.add_paragraph()


def split_table_row(line: str) -> list[str]:
    content = line.strip().strip("|")
    cells: list[str] = []
    current: list[str] = []
    escape = False

    for char in content:
        if escape:
            current.append(char)
            escape = False
            continue

        if char == "\\":
            escape = True
            current.append(char)
            continue

        if char == "|":
            cells.append("".join(current).strip())
            current = []
            continue

        current.append(char)

    cells.append("".join(current).strip())
    return cells


def is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:\-|\t]+\|?", line.strip()))


def skip_whitespace(text: str, index: int) -> int:
    while index < len(text) and text[index].isspace():
        index += 1
    return index


def find_matching_brace(text: str, start: int) -> int:
    depth = 0
    in_string = False
    quote_char = ""
    escape = False

    for index in range(start, len(text)):
        char = text[index]
        if escape:
            escape = False
            continue

        if char == "\\":
            escape = True
            continue

        if in_string:
            if char == quote_char:
                in_string = False
            continue

        if char in ("'", '"'):
            in_string = True
            quote_char = char
            continue

        if char == "{":
            depth += 1
            continue

        if char == "}":
            depth -= 1
            if depth == 0:
                return index

    return -1


def is_field_start(text: str, index: int) -> bool:
    if index >= len(text) or not (text[index].isalpha() or text[index] == "_"):
        return False

    match = re.match(r"[A-Za-z_][A-Za-z0-9_]*", text[index:])
    if not match:
        return False

    next_index = skip_whitespace(text, index + len(match.group(0)))
    if next_index >= len(text):
        return True

    return (
        text[next_index] in (":", ",", "}")
        or text.startswith("//", next_index)
    )


def find_comment_boundary(text: str, start: int) -> int:
    depth = 0
    in_string = False
    quote_char = ""
    escape = False
    index = start

    while index < len(text):
        char = text[index]
        if escape:
            escape = False
            index += 1
            continue

        if char == "\\":
            escape = True
            index += 1
            continue

        if in_string:
            if char == quote_char:
                in_string = False
            index += 1
            continue

        if char in ("'", '"'):
            in_string = True
            quote_char = char
            index += 1
            continue

        if char == "{":
            depth += 1
            index += 1
            continue

        if char == "}":
            if depth == 0:
                return index
            depth -= 1
            index += 1
            continue

        if depth == 0 and is_field_start(text, index):
            return index

        index += 1

    return len(text)


def parse_scalar_value(text: str, index: int) -> tuple[str, int]:
    value_start = index
    depth = 0
    in_string = False
    quote_char = ""
    escape = False

    while index < len(text):
        char = text[index]
        if escape:
            escape = False
            index += 1
            continue

        if char == "\\":
            escape = True
            index += 1
            continue

        if in_string:
            if char == quote_char:
                in_string = False
            index += 1
            continue

        if char in ("'", '"'):
            in_string = True
            quote_char = char
            index += 1
            continue

        if char == "{":
            depth += 1
            index += 1
            continue

        if char == "}":
            if depth == 0:
                break
            depth -= 1
            index += 1
            continue

        if depth == 0 and char == ",":
            break

        if depth == 0 and text.startswith("//", index):
            break

        index += 1

    return text[value_start:index].strip(), index


def parse_json_like_object(text: str) -> list[JsonLikeField]:
    stripped = text.strip()
    if not (stripped.startswith("{") and stripped.endswith("}")):
        return []

    inner = stripped[1:-1]
    fields: list[JsonLikeField] = []
    index = 0

    while index < len(inner):
        index = skip_whitespace(inner, index)
        while index < len(inner) and inner[index] == ",":
            index += 1
            index = skip_whitespace(inner, index)

        if index >= len(inner):
            break

        key_match = re.match(r"[A-Za-z_][A-Za-z0-9_]*", inner[index:])
        if not key_match:
            index += 1
            continue

        key = key_match.group(0)
        index += len(key)
        index = skip_whitespace(inner, index)

        value: str | None = None
        children: list[JsonLikeField] | None = None
        comment: str | None = None

        if index < len(inner) and inner[index] == ":":
            index += 1
            index = skip_whitespace(inner, index)

            if index < len(inner) and inner[index] == "{":
                end_brace = find_matching_brace(inner, index)
                if end_brace == -1:
                    value = inner[index:].strip()
                    index = len(inner)
                else:
                    children = parse_json_like_object(inner[index:end_brace + 1])
                    index = end_brace + 1
            else:
                value, index = parse_scalar_value(inner, index)

        index = skip_whitespace(inner, index)
        if index < len(inner) and inner[index] == ",":
            index += 1
            index = skip_whitespace(inner, index)

        if index < len(inner) and inner.startswith("//", index):
            comment_start = index + 2
            comment_end = find_comment_boundary(inner, comment_start)
            comment = inner[comment_start:comment_end].strip()
            index = comment_end

        fields.append(JsonLikeField(key=key, value=value, children=children, comment=comment))

    return fields


def fallback_multiline_json_like(value: str) -> list[str]:
    stripped = value.strip()
    if not (stripped.startswith("{") and stripped.endswith("}")):
        return [stripped]

    normalized = stripped[1:-1].strip()
    normalized = normalized.replace(" request: {", ", request: {")
    normalized = normalized.replace(" response: {", ", response: {")
    normalized = normalized.replace(" }, response:", " },, response:")
    parts = [part.strip() for part in normalized.split(",") if part.strip()]

    lines = ["{"]
    for part in parts:
        lines.append(f"  {part}")
    lines.append("}")
    return lines


def render_json_like_fields(fields: list[JsonLikeField], indent: int = 0) -> list[str]:
    lines = [(" " * indent) + "{"]

    for field_index, field in enumerate(fields):
        is_last = field_index == len(fields) - 1
        prefix = " " * (indent + 2)

        if field.children is not None:
            lines.append(prefix + f"{field.key}: {{")
            nested_lines = render_json_like_fields(field.children, indent + 2)
            lines.extend(nested_lines[1:-1])
            closing_line = prefix + "}"
            if not is_last:
                closing_line += ","
            if field.comment:
                closing_line += f" // {field.comment}"
            lines.append(closing_line)
            continue

        line = prefix + field.key
        if field.value:
            line += f": {field.value}"
        if not is_last:
            line += ","
        if field.comment:
            line += f" // {field.comment}"
        lines.append(line)

    lines.append((" " * indent) + "}")
    return lines


def format_json_like_block(value: str, indent: int = 0) -> list[str]:
    stripped = value.strip().strip("`")
    if not (stripped.startswith("{") and stripped.endswith("}")):
        return [(" " * indent) + stripped]

    fields = parse_json_like_object(stripped)
    if not fields:
        return fallback_multiline_json_like(stripped)
    return render_json_like_fields(fields, indent)


def format_table_cell_text(header: str, value: str) -> str:
    if header.strip() != "数据":
        return value

    stripped = value.strip().strip("`").replace("\\|", "|")
    lines = format_json_like_block(stripped)
    if len(lines) <= 2:
        lines = fallback_multiline_json_like(stripped)
    return "\n".join(lines)


def add_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if line.strip()]
    if len(rows) < 2:
        return

    headers = rows[0]
    header_len = len(headers)
    body_rows = rows[2:] if len(rows) >= 2 and is_table_separator(lines[1]) else rows[1:]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for index, value in enumerate(headers):
        header_cells[index].text = value
        set_cell_shading(header_cells[index], "D9EAF7")
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True

    for body in body_rows:
        normalized_body = body[:header_len] + [""] * max(0, header_len - len(body))
        row_cells = table.add_row().cells
        for index, value in enumerate(normalized_body):
            row_cells[index].text = format_table_cell_text(headers[index], value)

    document.add_paragraph()


def flush_paragraph_buffer(document: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    add_paragraph(document, " ".join(item.strip() for item in buffer if item.strip()))
    buffer.clear()


def build_docx() -> None:
    content = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    set_doc_defaults(document)

    paragraph_buffer: list[str] = []
    code_buffer: list[str] = []
    table_buffer: list[str] = []
    in_code_block = False

    for line in content:
        stripped = line.rstrip()

        if stripped.startswith("```"):
            flush_paragraph_buffer(document, paragraph_buffer)
            if in_code_block:
                add_code_block(document, code_buffer)
                code_buffer.clear()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(stripped)
            continue

        if stripped.startswith("|"):
            flush_paragraph_buffer(document, paragraph_buffer)
            table_buffer.append(stripped)
            continue

        if table_buffer:
            add_table(document, table_buffer)
            table_buffer.clear()

        if not stripped.strip():
            flush_paragraph_buffer(document, paragraph_buffer)
            continue

        if stripped.startswith("# "):
            flush_paragraph_buffer(document, paragraph_buffer)
            add_title(document, stripped[2:].strip())
            continue

        if stripped.startswith("## "):
            flush_paragraph_buffer(document, paragraph_buffer)
            add_heading(document, 1, stripped[3:].strip())
            continue

        if stripped.startswith("### "):
            flush_paragraph_buffer(document, paragraph_buffer)
            add_heading(document, 2, stripped[4:].strip())
            continue

        bullet_match = re.match(r"^-\s+(.*)", stripped)
        if bullet_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            add_list_item(document, bullet_match.group(1).strip(), ordered=False)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if ordered_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            add_list_item(document, ordered_match.group(1).strip(), ordered=True)
            continue

        paragraph_buffer.append(stripped)

    if table_buffer:
        add_table(document, table_buffer)
    flush_paragraph_buffer(document, paragraph_buffer)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)


if __name__ == "__main__":
    build_docx()
